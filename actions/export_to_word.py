"""
title: Export to Word (APA 7th – Merged)
author: Fu-Jie (original); Marios Adamidis / FORTHought Lab (APA/Greek enhancements)
author_url: https://github.com/Fu-Jie
funding_url: https://github.com/Fu-Jie/awesome-openwebui
version: 1.0.0
icon_url: data:image/svg+xml;base64,PHN2ZwogIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyIKICB3aWR0aD0iMjQiCiAgaGVpZ2h0PSIyNCIKICB2aWV3Qm94PSIwIDAgMjQgMjQiCiAgZmlsbD0ibm9uZSIKICBzdHJva2U9ImN1cnJlbnRDb2xvciIKICBzdHJva2Utd2lkdGg9IjIiCiAgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIgogIHN0cm9rZS1saW5lam9pbj0icm91bmQiCj4KICA8cGF0aCBkPSJNNiAyMmEyIDIgMCAwIDEtMi0yVjRhMiAyIDAgMCAxIDItMmg4YTIuNCAyLjQgMCAwIDEgMS43MDQuNzA2bDMuNTg4IDMuNTg4QTIuNCAyLjQgMCAwIDEgMjAgOHYxMmEyIDIgMCAwIDEtMiAyeiIgLz4KICA8cGF0aCBkPSJNMTQgMnY1YTEgMSAwIDAgMCAxIDFoNSIgLz4KICA8cGF0aCBkPSJNMTAgOUg4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxM0g4IiAvPgogIDxwYXRoIGQ9Ik0xNiAxN0g4IiAvPgo8L3N2Zz4K
requirements: python-docx, Pygments, latex2mathml, mathml2omml
description: Export to Word with strict APA 7th Edition styling, Greek/English i18n, LaTeX math, Mermaid diagrams, S3 support, debug logging, and robust image embedding (merged upstream v0.4.4).
notes: Merges upstream v0.4.4 improvements (S3/boto3, debug logging, richer image resolution, robust UserValves, East Asian font fallback) into the APA/Greek fork.
"""

from __future__ import annotations

import re
import base64
import datetime
import time
import io
import asyncio
import logging
import hashlib
import struct
import os
import glob
import zlib
import binascii
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, Callable, Awaitable, Any, List, Tuple, Dict, cast
from urllib.parse import quote
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from open_webui.models.chats import Chats
from open_webui.models.users import Users
from open_webui.utils.chat import generate_chat_completion
from pydantic import BaseModel, Field

# Files are used to embed internal /api/v1/files/<id>/content images.
try:
    from open_webui.models.files import Files  # type: ignore
except Exception:  # pragma: no cover
    Files = None

# Pygments for syntax highlighting
try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token

    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False

# Math Conversion Imports
try:
    from latex2mathml.converter import convert as latex_to_mathml
    import mathml2omml

    LATEX_MATH_AVAILABLE = True
except ImportError:
    LATEX_MATH_AVAILABLE = False

# boto3 for S3 direct access (ported from upstream v0.4.4)
try:
    import boto3
    from botocore.config import Config as BotoConfig

    BOTO3_AVAILABLE = True
except ImportError:
    BOTO3_AVAILABLE = False


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

_AUTO_URL_RE = re.compile(r"(?:https?://|www\.)[^\s<>()]+")
_DATA_IMAGE_URL_RE = re.compile(
    r"^data:(?P<mime>image/[a-z0-9.+-]+)\s*;\s*base64\s*,\s*(?P<b64>.*)$",
    re.IGNORECASE | re.DOTALL,
)
_OWUI_API_FILE_ID_RE = re.compile(
    r"/api/v1/files/(?P<id>[A-Za-z0-9-]+)(?:/content)?(?:[/?#]|$)",
    re.IGNORECASE,
)
_CURRENCY_NUMBER_RE = re.compile(r"^\d[\d,]*(?:\.\d+)?$")

_TRANSPARENT_1PX_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQImWNgYGBgAAAABQABDQottAAAAABJRU5ErkJggg=="
)

_ASVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
nsmap.setdefault("asvg", _ASVG_NS)

_ALL_DETAILS_RE = re.compile(
    r"<details\b[^>]*>.*?</details\s*>", re.IGNORECASE | re.DOTALL
)
_THINK_RE = re.compile(r"<think\b[^>]*>.*?</think\s*>", re.IGNORECASE | re.DOTALL)
_ANALYSIS_RE = re.compile(
    r"<analysis\b[^>]*>.*?</analysis\s*>", re.IGNORECASE | re.DOTALL
)


@dataclass(frozen=True)
class _CitationRef:
    idx: int
    anchor: str
    title: str
    url: Optional[str]
    source_id: str


class Action:
    # ── i18n: English + Greek only ──────────────────────────────────────
    _I18N_MESSAGES: Dict[str, Dict[str, str]] = {
        "en": {
            "converting": "Converting to Word document...",
            "exported": "Word document exported",
            "success": "Successfully exported to {filename}",
            "error_no_content": "No content found to export!",
            "error_export": "Error exporting Word document: {error}",
            "export_failed": "Export failed: {error}",
            "figure_prefix": "Figure",
            "references": "References",
        },
        "el": {
            "converting": "Μετατροπή σε έγγραφο Word...",
            "exported": "Το έγγραφο Word εξήχθη",
            "success": "Επιτυχής εξαγωγή στο {filename}",
            "error_no_content": "Δεν βρέθηκε περιεχόμενο για εξαγωγή!",
            "error_export": "Σφάλμα κατά την εξαγωγή: {error}",
            "export_failed": "Η εξαγωγή απέτυχε: {error}",
            "figure_prefix": "Σχήμα",
            "references": "Βιβλιογραφία",
        },
    }

    # ── Valves ───────────────────────────────────────────────────────────
    class Valves(BaseModel):
        TITLE_SOURCE: str = Field(
            default="chat_title",
            description="Title Source: 'chat_title', 'ai_generated', 'markdown_title'",
        )
        SHOW_STATUS: bool = Field(
            default=True,
            description="Whether to show operation status updates.",
        )
        SHOW_DEBUG_LOG: bool = Field(
            default=False,
            description="Whether to print debug logs in the browser console.",
        )

        MAX_EMBED_IMAGE_MB: int = Field(
            default=20,
            description="Maximum image size to embed into DOCX (MB).",
        )

        # Font configuration – APA defaults
        FONT_LATIN: str = Field(
            default="Times New Roman",
            description="Primary font (APA standard: Times New Roman)",
        )
        FONT_ASIAN: str = Field(
            default="Times New Roman",
            description="Fallback font for East Asian glyphs",
        )
        FONT_CODE: str = Field(
            default="Consolas",
            description="Font for code blocks (e.g., 'Consolas', 'Courier New')",
        )

        # Title alignment (upstream feature)
        TITLE_ALIGNMENT: str = Field(
            default="center",
            description="Title alignment: 'left', 'center', or 'right'",
        )

        # Table styling
        TABLE_HEADER_COLOR: str = Field(
            default="E8E8E8",
            description="Table header background color (hex, without #)",
        )
        TABLE_ZEBRA_COLOR: str = Field(
            default="F8F8F8",
            description="Table zebra stripe background color (hex, without #)",
        )

        MERMAID_JS_URL: str = Field(
            default="https://cdn.jsdelivr.net/npm/mermaid@11.12.2/dist/mermaid.min.js",
            description="Mermaid JS CDN URL",
        )
        MERMAID_JSZIP_URL: str = Field(
            default="https://cdnjs.cloudflare.com/ajax/libs/jszip/3.10.1/jszip.min.js",
            description="JSZip CDN URL",
        )
        MERMAID_PNG_SCALE: float = Field(
            default=3.0,
            description="PNG render resolution multiplier",
        )
        MERMAID_DISPLAY_SCALE: float = Field(
            default=1.0,
            description="Diagram width relative to available page width",
        )
        MERMAID_OPTIMIZE_LAYOUT: bool = Field(
            default=False,
            description="Optimize Mermaid layout (convert LR to TD)",
        )
        MERMAID_BACKGROUND: str = Field(
            default="",
            description="Mermaid background color (empty = transparent)",
        )

        MERMAID_CAPTIONS_ENABLE: bool = Field(
            default=True,
            description="Add figure captions under diagrams",
        )
        MERMAID_CAPTION_STYLE: str = Field(
            default="Caption",
            description="Word Style for captions",
        )
        MERMAID_CAPTION_PREFIX: str = Field(
            default="",
            description="Caption prefix (e.g. 'Σχήμα'). Empty = auto-detect.",
        )

        MATH_ENABLE: bool = Field(
            default=True,
            description="Enable LaTeX math block conversion",
        )
        MATH_INLINE_DOLLAR_ENABLE: bool = Field(
            default=True,
            description="Enable inline $...$ math conversion",
        )

        UI_LANGUAGE: str = Field(
            default="el",
            description="UI language for export messages. Options: 'en', 'el'",
        )

    class UserValves(BaseModel):
        TITLE_SOURCE: Optional[str] = Field(
            default=None,
            description="Title Source: 'chat_title', 'ai_generated', 'markdown_title'",
        )
        UI_LANGUAGE: Optional[str] = Field(
            default=None,
            description="UI language: 'en', 'el'",
        )
        FONT_LATIN: Optional[str] = Field(
            default=None,
            description="Primary font",
        )
        FONT_ASIAN: Optional[str] = Field(
            default=None,
            description="East Asian fallback font",
        )
        FONT_CODE: Optional[str] = Field(
            default=None,
            description="Code font",
        )
        TABLE_HEADER_COLOR: Optional[str] = Field(
            default=None,
            description="Table header background color (hex)",
        )
        TABLE_ZEBRA_COLOR: Optional[str] = Field(
            default=None,
            description="Table zebra stripe color (hex)",
        )
        MERMAID_PNG_SCALE: Optional[float] = Field(
            default=None,
            description="PNG render resolution multiplier",
        )
        MERMAID_DISPLAY_SCALE: Optional[float] = Field(
            default=None,
            description="Diagram width relative to page width",
        )
        MERMAID_OPTIMIZE_LAYOUT: Optional[bool] = Field(
            default=None,
            description="Optimize Mermaid layout (LR→TD)",
        )
        MERMAID_BACKGROUND: Optional[str] = Field(
            default=None,
            description="Mermaid background color",
        )
        MERMAID_CAPTIONS_ENABLE: Optional[bool] = Field(
            default=None,
            description="Add figure captions",
        )
        MATH_ENABLE: Optional[bool] = Field(
            default=None,
            description="Enable LaTeX math",
        )
        MATH_INLINE_DOLLAR_ENABLE: Optional[bool] = Field(
            default=None,
            description="Enable inline $…$ math",
        )

    # ── init ─────────────────────────────────────────────────────────────
    def __init__(self):
        self.valves = self.Valves()
        self._mermaid_figure_counter: int = 0
        self._mermaid_placeholder_counter: int = 0
        self._caption_style_name: Optional[str] = None
        self._citation_anchor_by_index: Dict[int, str] = {}
        self._citation_refs: List[_CitationRef] = []
        self._bookmark_id_counter: int = 1
        self._active_doc: Optional[Document] = None
        self._user_lang: str = "el"
        # Ported from upstream: API connection info for internal file fetching
        self._api_token: Optional[str] = None
        self._api_base_url: Optional[str] = None

    # ── i18n helpers ─────────────────────────────────────────────────────
    def _get_lang_key(self, user_language: str) -> str:
        lang = (user_language or "el").lower().split("-")[0]
        if lang in ("el", "gr", "greek"):
            return "el"
        return lang if lang in self._I18N_MESSAGES else "el"

    def _get_msg(self, key: str, **kwargs) -> str:
        messages = self._I18N_MESSAGES.get(self._user_lang, self._I18N_MESSAGES["el"])
        msg = messages.get(key, self._I18N_MESSAGES["el"].get(key, key))
        if kwargs:
            try:
                return msg.format(**kwargs)
            except KeyError:
                return msg
        return msg

    # ── context helpers (ported from upstream) ───────────────────────────
    def _get_user_context(self, __user__: Optional[Dict[str, Any]]) -> Dict[str, str]:
        """Safely extracts user context information."""
        if isinstance(__user__, (list, tuple)):
            user_data = __user__[0] if __user__ else {}
        elif isinstance(__user__, dict):
            user_data = __user__
        else:
            user_data = {}

        return {
            "user_id": user_data.get("id", "unknown_user"),
            "user_name": user_data.get("name", "User"),
            "user_language": user_data.get("language", "en-US"),
        }

    def _get_chat_context(
        self, body: dict, __metadata__: Optional[dict] = None
    ) -> Dict[str, str]:
        """Unified extraction of chat context (chat_id, message_id)."""
        chat_id = ""
        message_id = ""

        if isinstance(body, dict):
            chat_id = body.get("chat_id", "")
            message_id = body.get("id", "")
            if not chat_id or not message_id:
                body_metadata = body.get("metadata", {})
                if isinstance(body_metadata, dict):
                    if not chat_id:
                        chat_id = body_metadata.get("chat_id", "")
                    if not message_id:
                        message_id = body_metadata.get("message_id", "")

        if __metadata__ and isinstance(__metadata__, dict):
            if not chat_id:
                chat_id = __metadata__.get("chat_id", "")
            if not message_id:
                message_id = __metadata__.get("message_id", "")

        return {
            "chat_id": str(chat_id).strip(),
            "message_id": str(message_id).strip(),
        }

    # ── emitter helpers (ported from upstream) ───────────────────────────
    async def _emit_status(
        self,
        emitter: Optional[Callable[[Any], Awaitable[None]]],
        description: str,
        done: bool = False,
    ):
        if self.valves.SHOW_STATUS and emitter:
            await emitter(
                {"type": "status", "data": {"description": description, "done": done}}
            )

    async def _emit_notification(
        self,
        emitter: Optional[Callable[[Any], Awaitable[None]]],
        content: str,
        ntype: str = "info",
    ):
        if emitter:
            await emitter(
                {"type": "notification", "data": {"type": ntype, "content": content}}
            )

    async def _emit_debug_log(self, emitter, title: str, data: dict):
        """Print structured debug logs in the browser console."""
        if not self.valves.SHOW_DEBUG_LOG or not emitter:
            return
        try:
            import json

            js_code = f"""
                (async function() {{
                    console.group("🛠️ {title}");
                    console.log({json.dumps(data, ensure_ascii=False)});
                    console.groupEnd();
                }})();
            """
            await emitter({"type": "execute", "data": {"code": js_code}})
        except Exception as e:
            print(f"Error emitting debug log: {e}")

    # ── main action ──────────────────────────────────────────────────────
    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__: Optional[Callable[[Any], Awaitable[None]]] = None,
        __metadata__: Optional[dict] = None,
        __request__: Optional[Any] = None,
    ):
        logger.info(f"action:{__name__}")

        # Parse user info
        user_ctx = self._get_user_context(__user__)
        user_name = user_ctx["user_name"]
        user_id = user_ctx["user_id"]

        # ── Robust UserValves handling (upstream fix: exclude_unset) ─────
        if __user__:
            raw_valves = (
                __user__[0].get("valves", {})
                if isinstance(__user__, (list, tuple)) and __user__
                else __user__.get("valves", {}) if isinstance(__user__, dict) else {}
            )
            if isinstance(raw_valves, self.UserValves):
                user_valves = raw_valves
            elif isinstance(raw_valves, dict):
                user_valves = self.UserValves(**raw_valves)
            else:
                user_valves = None

            if user_valves:
                for key, value in user_valves.model_dump(exclude_unset=True).items():
                    if hasattr(self.valves, key) and value is not None:
                        setattr(self.valves, key, value)

        self._user_lang = self._get_lang_key(self.valves.UI_LANGUAGE)

        # ── Extract API connection info (ported from upstream) ───────────
        def _get_default_base_url() -> str:
            port = os.environ.get("PORT") or "8080"
            return f"http://localhost:{port}"

        if __request__:
            try:
                self._api_token = __request__.headers.get("Authorization")
                self._api_base_url = str(__request__.base_url).rstrip("/")
            except Exception:
                self._api_token = None
                self._api_base_url = _get_default_base_url()
        else:
            self._api_token = None
            self._api_base_url = _get_default_base_url()

        if __event_emitter__:
            last_assistant_message = body["messages"][-1]

            await self._emit_status(
                __event_emitter__, self._get_msg("converting"), done=False
            )

            try:
                message_content = last_assistant_message["content"]
                if isinstance(message_content, str):
                    # Debug logging (upstream feature)
                    if __event_emitter__ and self.valves.SHOW_DEBUG_LOG:
                        debug_data = {}
                        for name, regex in [
                            ("Details Block", _ALL_DETAILS_RE),
                            ("Think Block", _THINK_RE),
                            ("Analysis Block", _ANALYSIS_RE),
                        ]:
                            matches = regex.findall(message_content)
                            if matches:
                                debug_data[name] = [
                                    (m[:200] + "...") if len(m) > 200 else m
                                    for m in matches
                                ]
                        if debug_data:
                            await self._emit_debug_log(
                                __event_emitter__,
                                "Context Stripping Analysis",
                                debug_data,
                            )
                        await self._emit_debug_log(
                            __event_emitter__,
                            "Font Configuration",
                            {
                                "Latin Font": self.valves.FONT_LATIN,
                                "Asian Font": self.valves.FONT_ASIAN,
                                "Code Font": self.valves.FONT_CODE,
                            },
                        )

                    message_content = self._strip_reasoning_blocks(message_content)
                    message_content = self._strip_emojis(message_content)

                if not message_content or not message_content.strip():
                    await self._emit_notification(
                        __event_emitter__,
                        self._get_msg("error_no_content"),
                        "error",
                    )
                    return

                # ── Generate filename ────────────────────────────────────
                title = ""
                chat_ctx = self._get_chat_context(body, __metadata__)
                chat_id = chat_ctx["chat_id"]

                chat_title = ""
                if chat_id:
                    chat_title = await self.fetch_chat_title(chat_id, user_id)
                    # OWUI auto-generates chat titles with emojis (📈, 🔬, etc.)
                    # Strip them early so they don't leak into heading or filename
                    chat_title = self._remove_emojis(chat_title).strip()

                if (
                    self.valves.TITLE_SOURCE.strip() == "chat_title"
                    or not self.valves.TITLE_SOURCE.strip()
                ):
                    title = chat_title
                elif self.valves.TITLE_SOURCE.strip() == "markdown_title":
                    title = self.extract_title(message_content)
                elif self.valves.TITLE_SOURCE.strip() == "ai_generated":
                    title = await self.generate_title_using_ai(
                        body, message_content, user_id, __request__
                    )

                if not title:
                    if self.valves.TITLE_SOURCE.strip() != "chat_title" and chat_title:
                        title = chat_title
                    elif self.valves.TITLE_SOURCE.strip() != "markdown_title":
                        extracted = self.extract_title(message_content)
                        if extracted:
                            title = extracted

                current_datetime = datetime.datetime.now()
                formatted_date = current_datetime.strftime("%Y%m%d")

                cleaned_title = self.clean_filename(title) if title else ""
                if cleaned_title:
                    filename = f"{cleaned_title}.docx"
                else:
                    clean_user = self.clean_filename(user_name)
                    filename = f"{clean_user}_{formatted_date}.docx"

                js_filename = filename.replace("\\", "\\\\").replace('"', '\\"')

                top_heading = chat_title or title or ""
                # OWUI auto-generates chat titles with emojis (e.g. 📈, 🔬)
                # Strip them for a clean document heading
                top_heading = self._remove_emojis(top_heading).strip()

                # ── Create Word document ─────────────────────────────────
                has_h1 = bool(re.search(r"^#\s+.+$", message_content, re.MULTILINE))
                sources = (
                    last_assistant_message.get("sources") or body.get("sources") or []
                )
                doc = await self.markdown_to_docx(
                    message_content,
                    top_heading=top_heading,
                    has_h1=has_h1,
                    sources=sources,
                    event_emitter=__event_emitter__,
                )

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                file_content = doc_buffer.read()
                base64_blob = base64.b64encode(file_content).decode("utf-8")

                # ── Trigger file download (client-side Mermaid pipeline) ─
                if __event_call__:
                    await __event_call__(
                        {
                            "type": "execute",
                            "data": {
                                "code": f"""
                                (async function() {{
                                    const base64Data = "{base64_blob}";
                                    const filename = "{js_filename}";
                                    const mermaidUrl = "{self.valves.MERMAID_JS_URL}";
                                    const jszipUrl = "{self.valves.MERMAID_JSZIP_URL}";
                                    const pngScale = {float(self.valves.MERMAID_PNG_SCALE)};
                                    const displayScale = {float(self.valves.MERMAID_DISPLAY_SCALE)};
                                    const bgRaw = "{(self.valves.MERMAID_BACKGROUND or '').strip()}";
                                    const bg = (bgRaw || "").trim();
                                    const bgFill = (bg && bg.toLowerCase() !== "transparent") ? bg : "";
                                    const themeBackground = bgFill || "transparent";

                                    function downloadBlob(blob, filename) {{
                                        const url = URL.createObjectURL(blob);
                                        const a = document.createElement("a");
                                        a.style.display = "none";
                                        a.href = url;
                                        a.download = filename;
                                        document.body.appendChild(a);
                                        a.click();
                                        URL.revokeObjectURL(url);
                                        document.body.removeChild(a);
                                    }}

                                    async function loadScript(url, globalName) {{
                                        if (globalName && window[globalName]) return;
                                        await new Promise((resolve, reject) => {{
                                            const script = document.createElement("script");
                                            script.src = url;
                                            script.onload = resolve;
                                            script.onerror = reject;
                                            document.head.appendChild(script);
                                        }});
                                    }}

                                    function decodeBase64ToUint8Array(b64) {{
                                        const binary = atob(b64);
                                        const bytes = new Uint8Array(binary.length);
                                        for (let i = 0; i < binary.length; i++) bytes[i] = binary.charCodeAt(i);
                                        return bytes;
                                    }}

                                    function parseViewBox(vb) {{
                                        if (!vb) return null;
                                        const parts = vb.trim().split(/\\s+/).map(Number);
                                        if (parts.length !== 4 || parts.some((n) => !isFinite(n))) return null;
                                        return {{ minX: parts[0], minY: parts[1], width: parts[2], height: parts[3] }};
                                    }}

                                    function normalizeSvgForWord(svgText) {{
                                        const parser = new DOMParser();
                                        const doc = parser.parseFromString(svgText, "image/svg+xml");
                                        const svgEl = doc.documentElement;
                                        if (!svgEl || svgEl.tagName.toLowerCase() !== "svg") return svgText;

                                        const vb0 = parseViewBox(svgEl.getAttribute("viewBox"));
                                        if (vb0 && vb0.width > 0 && vb0.height > 0) {{
                                            const minDim = Math.min(vb0.width, vb0.height);
                                            let pad = Math.max(8.0, minDim * 0.02);
                                            pad = Math.min(pad, 24.0);
                                            const vb = {{
                                                minX: vb0.minX - pad,
                                                minY: vb0.minY - pad,
                                                width: vb0.width + 2 * pad,
                                                height: vb0.height + 2 * pad,
                                            }};
                                            svgEl.setAttribute("viewBox", `${{vb.minX}} ${{vb.minY}} ${{vb.width}} ${{vb.height}}`);
                                        }}

                                        const vb = parseViewBox(svgEl.getAttribute("viewBox"));
                                        const widthAttr = (svgEl.getAttribute("width") || "").trim();
                                        const heightAttr = (svgEl.getAttribute("height") || "").trim();
                                        const widthPct = widthAttr.endsWith("%");
                                        const heightPct = heightAttr.endsWith("%");
                                        if (vb && vb.width > 0 && vb.height > 0 && (!widthAttr || !heightAttr || widthPct || heightPct)) {{
                                            svgEl.setAttribute("width", `${{vb.width}}`);
                                            svgEl.setAttribute("height", `${{vb.height}}`);
                                        }}

                                        svgEl.removeAttribute("style");
                                        svgEl.setAttribute("preserveAspectRatio", "xMidYMid meet");
                                        svgEl.setAttribute("overflow", "visible");

                                        const removeNode = (n) => {{
                                            try {{ n && n.parentNode && n.parentNode.removeChild(n); }} catch (_e) {{}}
                                        }};

                                        svgEl
                                            .querySelectorAll('rect[data-owui-bg="1"], rect.background, rect[class~="background"], rect#background')
                                            .forEach(removeNode);
                                        try {{
                                            const isWhiteish = (fill) => {{
                                                const f = (fill || "").trim().toLowerCase();
                                                return (
                                                    f === "white" ||
                                                    f === "#fff" ||
                                                    f === "#ffffff" ||
                                                    f === "rgb(255,255,255)" ||
                                                    f === "rgb(255, 255, 255)"
                                                );
                                            }};
                                            const nearly = (a, b) => Math.abs(a - b) <= 1e-3;
                                            const rectMatches = (r, box) => {{
                                                if (!box) return false;
                                                const x = parseFloat(r.getAttribute("x") || "0");
                                                const y = parseFloat(r.getAttribute("y") || "0");
                                                const w = parseFloat(r.getAttribute("width") || "");
                                                const h = parseFloat(r.getAttribute("height") || "");
                                                if (!isFinite(x) || !isFinite(y) || !isFinite(w) || !isFinite(h)) return false;
                                                return (
                                                    nearly(x, box.minX) &&
                                                    nearly(y, box.minY) &&
                                                    nearly(w, box.width) &&
                                                    nearly(h, box.height)
                                                );
                                            }};
                                            const vbNow = parseViewBox(svgEl.getAttribute("viewBox"));
                                            svgEl.querySelectorAll("rect[fill]").forEach((r) => {{
                                                const fill = r.getAttribute("fill");
                                                if (!isWhiteish(fill)) return;
                                                if (rectMatches(r, vb0) || rectMatches(r, vbNow)) removeNode(r);
                                            }});
                                        }} catch (_e) {{}}
                                        try {{
                                            const vbCanvas = parseViewBox(svgEl.getAttribute("viewBox")) || vb0 || vb;
                                            if (vbCanvas) {{
                                                const existing = svgEl.querySelector('rect[data-owui-canvas="1"]');
                                                const rect = existing || doc.createElementNS("http://www.w3.org/2000/svg", "rect");
                                                rect.setAttribute("data-owui-canvas", "1");
                                                rect.setAttribute("x", `${{vbCanvas.minX}}`);
                                                rect.setAttribute("y", `${{vbCanvas.minY}}`);
                                                rect.setAttribute("width", `${{vbCanvas.width}}`);
                                                rect.setAttribute("height", `${{vbCanvas.height}}`);
                                                rect.setAttribute("fill", "#FFFFFF");
                                                rect.setAttribute("fill-opacity", "0.001");
                                                rect.setAttribute("stroke", "none");
                                                rect.setAttribute("stroke-opacity", "0");
                                                rect.setAttribute("pointer-events", "none");
                                                if (!existing) {{
                                                    const first = svgEl.firstChild;
                                                    svgEl.insertBefore(rect, first);
                                                }}
                                            }}
                                        }} catch (_e) {{}}

                                        return new XMLSerializer().serializeToString(svgEl);
                                    }}

                                    function getMaxWidthEmu(xmlDoc) {{
                                        try {{
                                            const sects = xmlDoc.getElementsByTagName("w:sectPr");
                                            const sect = sects && sects.length ? sects[sects.length - 1] : null;
                                            if (!sect) return 5486400;
                                            const pgSz = sect.getElementsByTagName("w:pgSz")[0];
                                            const pgMar = sect.getElementsByTagName("w:pgMar")[0];
                                            if (!pgSz || !pgMar) return 5486400;
                                            const pageW = parseInt(pgSz.getAttribute("w:w") || "", 10);
                                            const left = parseInt(pgMar.getAttribute("w:left") || "", 10);
                                            const right = parseInt(pgMar.getAttribute("w:right") || "", 10);
                                            if (!isFinite(pageW) || !isFinite(left) || !isFinite(right)) return 5486400;
                                            const twips = Math.max(1, pageW - left - right);
                                            return Math.round(twips * 635);
                                        }} catch (_e) {{
                                            return 5486400;
                                        }}
                                    }}

                                    function getChildByTag(parent, tag) {{
                                        const nodes = parent.getElementsByTagName(tag);
                                        return nodes && nodes.length ? nodes[0] : null;
                                    }}

                                    try {{
                                        await loadScript(jszipUrl, "JSZip");
                                        await loadScript(mermaidUrl, "mermaid");

                                        try {{
                                            window.mermaid.initialize({{
                                                startOnLoad: false,
                                                theme: "default",
                                                themeVariables: {{
                                                    background: themeBackground,
                                                    fontFamily: "Times New Roman, Arial, sans-serif",
                                                    fontSize: "12pt",
                                                }},
                                                fontFamily: "Times New Roman, Arial, sans-serif",
                                                securityLevel: "strict",
                                                flowchart: {{ htmlLabels: false }},
                                            }});
                                        }} catch (_e) {{}}

                                        const bytes = decodeBase64ToUint8Array(base64Data);
                                        const zip = new window.JSZip();
                                        await zip.loadAsync(bytes);

                                        const docXml = await zip.file("word/document.xml").async("string");
                                        const relsXml = await zip.file("word/_rels/document.xml.rels").async("string");
                                        const parser = new DOMParser();
                                        const xmlDoc = parser.parseFromString(docXml, "application/xml");
                                        const relsDoc = parser.parseFromString(relsXml, "application/xml");

                                        const rels = relsDoc.getElementsByTagName("Relationship");
                                        const rIdToTarget = {{}};
                                        for (let i = 0; i < rels.length; i++) {{
                                            const rel = rels[i];
                                            const id = rel.getAttribute("Id");
                                            const target = rel.getAttribute("Target");
                                            if (id && target) rIdToTarget[id] = target;
                                        }}

                                        const maxWidthEmu = getMaxWidthEmu(xmlDoc);
                                        const maxWidthEmuScaled = Math.max(1, Math.round(maxWidthEmu * Math.min(1.0, Math.max(0.1, displayScale || 1.0))));

                                        const drawings = xmlDoc.getElementsByTagName("w:drawing");
                                        const placeholders = [];

                                        for (let i = 0; i < drawings.length; i++) {{
                                            const drawing = drawings[i];
                                            const docPr = getChildByTag(drawing, "wp:docPr");
                                            if (!docPr) continue;
                                            const descr = docPr.getAttribute("descr") || "";
                                            if (!descr.startsWith("MERMAID_SRC:")) continue;
                                            const encoded = descr.substring("MERMAID_SRC:".length);
                                            const code = decodeURIComponent(encoded);

                                            const blip = getChildByTag(drawing, "a:blip");
                                            const ridPng = blip ? blip.getAttribute("r:embed") : null;
                                            const svgBlip = getChildByTag(drawing, "asvg:svgBlip");
                                            const ridSvg = svgBlip ? svgBlip.getAttribute("r:embed") : null;

                                            const container = getChildByTag(drawing, "wp:inline") || getChildByTag(drawing, "wp:anchor");
                                            const extent = container ? getChildByTag(container, "wp:extent") : null;

                                            const xfrm = getChildByTag(drawing, "a:xfrm");
                                            const xfrmExt = xfrm ? getChildByTag(xfrm, "a:ext") : null;

                                            placeholders.push({{ code, ridPng, ridSvg, extent, xfrmExt, svgBlip }});
                                        }}

                                        if (!placeholders.length) {{
                                            const blob = new Blob([bytes], {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
                                            downloadBlob(blob, filename);
                                            return;
                                        }}

                                        const renderResults = [];
                                        for (let i = 0; i < placeholders.length; i++) {{
                                            const item = placeholders[i];
                                            try {{
                                                const id = "owui-mermaid-" + i;
                                                const rendered = await window.mermaid.render(id, item.code);
                                                let svgText = rendered && rendered.svg ? rendered.svg : rendered;
                                                if (!svgText || typeof svgText !== "string") throw new Error("Mermaid returned empty SVG");

                                                svgText = normalizeSvgForWord(svgText);
                                                const hasForeignObject = /<foreignObject\\b/i.test(svgText);
                                                if (hasForeignObject && item.svgBlip) {{
                                                    try {{ item.svgBlip.parentNode && item.svgBlip.parentNode.removeChild(item.svgBlip); }} catch (_e) {{}}
                                                    item.ridSvg = null;
                                                }}

                                                const svgDoc = new DOMParser().parseFromString(svgText, "image/svg+xml");
                                                const svgEl = svgDoc.documentElement;
                                                const vb = parseViewBox(svgEl && svgEl.getAttribute ? svgEl.getAttribute("viewBox") : null);
                                                const ratio = vb && vb.width > 0 && vb.height > 0 ? (vb.width / vb.height) : (4/3);

                                                const widthEmu = maxWidthEmuScaled;
                                                const heightEmu = Math.max(1, Math.round(widthEmu / ratio));

                                                renderResults.push({{ item, svgText, widthEmu, heightEmu, success: true }});
                                            }} catch (err) {{
                                                console.error("Mermaid render failed for block", i, err);
                                                renderResults.push({{ item, svgText: null, widthEmu: 0, heightEmu: 0, success: false }});
                                            }}
                                        }}

                                        async function svgToPng(svgText, targetWidthPx, targetHeightPx) {{
                                            const canvas = document.createElement("canvas");
                                            const ctx = canvas.getContext("2d");
                                            const scale = Math.max(1.0, pngScale || 1.0);
                                            canvas.width = Math.round(targetWidthPx * scale);
                                            canvas.height = Math.round(targetHeightPx * scale);
                                            ctx.setTransform(1, 0, 0, 1, 0, 0);
                                            if (bgFill) {{
                                                ctx.fillStyle = bgFill;
                                                ctx.fillRect(0, 0, canvas.width, canvas.height);
                                            }}
                                            ctx.scale(scale, scale);

                                            const img = new Image();
                                            await new Promise((resolve, reject) => {{
                                                img.onload = resolve;
                                                img.onerror = reject;
                                                img.src = "data:image/svg+xml;base64," + btoa(unescape(encodeURIComponent(svgText)));
                                            }});

                                            ctx.drawImage(img, 0, 0, targetWidthPx, targetHeightPx);
                                            const pngDataUrl = canvas.toDataURL("image/png");
                                            return pngDataUrl.split(",")[1];
                                        }}

                                        const pngPromises = renderResults.map(async (result, i) => {{
                                            if (!result.success || !result.svgText) return null;
                                            const {{ item, widthEmu, heightEmu }} = result;
                                            if (!item.ridPng || !rIdToTarget[item.ridPng]) return null;

                                            const targetWidthPx = Math.max(1, Math.round(widthEmu / 9525));
                                            const targetHeightPx = Math.max(1, Math.round(heightEmu / 9525));

                                            try {{
                                                const pngBase64 = await svgToPng(result.svgText, targetWidthPx, targetHeightPx);
                                                return {{ index: i, pngBase64, path: "word/" + rIdToTarget[item.ridPng] }};
                                            }} catch (err) {{
                                                console.error("PNG conversion failed for block", i, err);
                                                return null;
                                            }}
                                        }});

                                        const pngResults = await Promise.all(pngPromises);

                                        for (let i = 0; i < renderResults.length; i++) {{
                                            const result = renderResults[i];
                                            if (!result.success) continue;

                                            const {{ item, svgText, widthEmu, heightEmu }} = result;

                                            if (item.extent) {{
                                                item.extent.setAttribute("cx", `${{widthEmu}}`);
                                                item.extent.setAttribute("cy", `${{heightEmu}}`);
                                            }}
                                            if (item.xfrmExt) {{
                                                item.xfrmExt.setAttribute("cx", `${{widthEmu}}`);
                                                item.xfrmExt.setAttribute("cy", `${{heightEmu}}`);
                                            }}

                                            if (item.ridSvg && rIdToTarget[item.ridSvg]) {{
                                                zip.file("word/" + rIdToTarget[item.ridSvg], svgText);
                                            }}
                                        }}

                                        for (const pngResult of pngResults) {{
                                            if (pngResult && pngResult.pngBase64) {{
                                                zip.file(pngResult.path, pngResult.pngBase64, {{ base64: true }});
                                            }}
                                        }}

                                        const newDocXml = new XMLSerializer().serializeToString(xmlDoc);
                                        zip.file("word/document.xml", newDocXml);

                                        const finalBlob = await zip.generateAsync({{
                                            type: "blob",
                                            compression: "DEFLATE",
                                            compressionOptions: {{ level: 6 }},
                                        }});
                                        downloadBlob(finalBlob, filename);
                                    }} catch (error) {{
                                        console.error("Export pipeline failed:", error);
                                        const bytes = decodeBase64ToUint8Array(base64Data);
                                        const blob = new Blob([bytes], {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
                                        downloadBlob(blob, filename);
                                    }}
                                }})();
                                """
                            },
                        }
                    )

                await self._emit_status(
                    __event_emitter__, self._get_msg("exported"), done=True
                )
                await self._emit_notification(
                    __event_emitter__,
                    self._get_msg("success", filename=filename),
                    "success",
                )

                return {"message": "Download triggered"}

            except Exception as e:
                logger.exception(f"Error exporting to Word: {str(e)}")
                await self._emit_status(
                    __event_emitter__,
                    self._get_msg("export_failed", error=str(e)),
                    done=True,
                )
                await self._emit_notification(
                    __event_emitter__,
                    self._get_msg("error_export", error=str(e)),
                    "error",
                )

    # ── title helpers ────────────────────────────────────────────────────
    async def generate_title_using_ai(
        self, body: dict, content: str, user_id: str, request: Any
    ) -> str:
        if not request:
            return ""
        try:
            user_obj = Users.get_user_by_id(user_id)
            model = body.get("model")
            payload = {
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a helpful assistant. Generate a short, concise title (max 10 words) for the following text. Do not use quotes. Only output the title.",
                    },
                    {"role": "user", "content": content[:2000]},
                ],
                "stream": False,
            }
            response = await generate_chat_completion(request, payload, user_obj)
            if response and "choices" in response:
                return response["choices"][0]["message"]["content"].strip()
        except Exception as e:
            logger.error(f"Error generating title: {e}")
        return ""

    def extract_title(self, content: str) -> str:
        lines = content.split("\n")
        for line in lines:
            match = re.match(r"^#{1,2}\s+(.+)$", line.strip())
            if match:
                return match.group(1).strip()
        return ""

    def extract_chat_id(self, body: dict, metadata: Optional[dict]) -> str:
        if isinstance(body, dict):
            chat_id = body.get("chat_id") or body.get("id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
            for key in ("chat", "conversation"):
                nested = body.get(key)
                if isinstance(nested, dict):
                    nested_id = nested.get("id") or nested.get("chat_id")
                    if isinstance(nested_id, str) and nested_id.strip():
                        return nested_id.strip()
        if isinstance(metadata, dict):
            chat_id = metadata.get("chat_id")
            if isinstance(chat_id, str) and chat_id.strip():
                return chat_id.strip()
        return ""

    async def fetch_chat_title(self, chat_id: str, user_id: str = "") -> str:
        if not chat_id:
            return ""

        def _load_chat():
            if user_id:
                chat = Chats.get_chat_by_id_and_user_id(id=chat_id, user_id=user_id)
                if chat:
                    return chat
            return Chats.get_chat_by_id(chat_id)

        try:
            chat = await asyncio.to_thread(_load_chat)
        except Exception as exc:
            logger.warning(f"Failed to load chat {chat_id}: {exc}")
            return ""
        if not chat:
            return ""
        data = getattr(chat, "chat", {}) or {}
        title = data.get("title") or getattr(chat, "title", "")
        return title.strip() if isinstance(title, str) else ""

    # ── filename / emoji helpers ─────────────────────────────────────────
    @staticmethod
    def _is_emoji_codepoint(codepoint: int) -> bool:
        return (
            0x1F000 <= codepoint <= 0x1FAFF
            or 0x1F1E6 <= codepoint <= 0x1F1FF
            or 0x2600 <= codepoint <= 0x26FF
            or 0x2700 <= codepoint <= 0x27BF
            or 0x2300 <= codepoint <= 0x23FF
            or 0x2B00 <= codepoint <= 0x2BFF
        )

    @staticmethod
    def _is_emoji_modifier(codepoint: int) -> bool:
        return (
            codepoint in (0x200D, 0xFE0E, 0xFE0F, 0x20E3)
            or 0x1F3FB <= codepoint <= 0x1F3FF
            or 0xE0020 <= codepoint <= 0xE007F
        )

    def _remove_emojis(self, text: str) -> str:
        if not isinstance(text, str):
            return ""
        return "".join(
            ch
            for ch in text
            if not (
                self._is_emoji_codepoint(ord(ch)) or self._is_emoji_modifier(ord(ch))
            )
        )

    def _strip_emojis(self, text: str) -> str:
        """Remove emojis from content body."""
        if not isinstance(text, str):
            return text
        return self._remove_emojis(text)

    def clean_filename(self, name: str) -> str:
        if not isinstance(name, str):
            return ""
        without_emoji = self._remove_emojis(name)
        cleaned = re.sub(r'[\\/*?:"<>|]', "", without_emoji)
        cleaned = re.sub(r"\s+", " ", cleaned).strip().strip(".")
        return cleaned[:50].strip()

    # ── image embedding (merged: upstream 6-step chain + APA glob fallback) ──
    def _max_embed_image_bytes(self) -> int:
        mb = getattr(self.valves, "MAX_EMBED_IMAGE_MB", 20)
        try:
            mb_i = int(mb)
        except Exception:
            mb_i = 20
        return max(1, mb_i) * 1024 * 1024

    def _extract_owui_api_file_id(self, url: str) -> Optional[str]:
        if not isinstance(url, str) or not url:
            return None
        m = _OWUI_API_FILE_ID_RE.search(url)
        if not m:
            return None
        fid = (m.group("id") or "").strip()
        return fid or None

    def _read_file_bytes_limited(self, path: Path, max_bytes: int) -> Optional[bytes]:
        try:
            if not path.exists():
                return None
            try:
                size = path.stat().st_size
                if size > max_bytes:
                    return None
            except Exception:
                pass
            with path.open("rb") as f:
                data = f.read(max_bytes + 1)
            if len(data) > max_bytes:
                return None
            return data
        except Exception:
            return None

    def _decode_base64_limited(self, b64: str, max_bytes: int) -> Optional[bytes]:
        if not isinstance(b64, str):
            return None
        s = re.sub(r"\s+", "", b64.strip())
        if not s:
            return None
        est = (len(s) * 3) // 4
        if est > max_bytes:
            return None
        pad = (-len(s)) % 4
        if pad:
            s = s + ("=" * pad)
        try:
            out = base64.b64decode(s, validate=False)
        except (binascii.Error, ValueError):
            return None
        if len(out) > max_bytes:
            return None
        return out

    def _image_bytes_from_data_url(self, url: str, max_bytes: int) -> Optional[bytes]:
        if not isinstance(url, str):
            return None
        m = _DATA_IMAGE_URL_RE.match(url.strip())
        if not m:
            return None
        b64 = m.group("b64") or ""
        return self._decode_base64_limited(b64, max_bytes)

    # ── S3 direct read (ported from upstream v0.4.4) ─────────────────────
    def _read_from_s3(self, s3_path: str, max_bytes: int) -> Optional[bytes]:
        """Read file directly from S3 using environment variables for credentials."""
        if not BOTO3_AVAILABLE:
            return None
        if not s3_path.startswith("s3://"):
            return None

        path_without_prefix = s3_path[5:]
        parts = path_without_prefix.split("/", 1)
        if len(parts) < 2:
            return None

        bucket = parts[0]
        key = parts[1]

        endpoint_url = os.environ.get("S3_ENDPOINT_URL")
        access_key = os.environ.get("S3_ACCESS_KEY_ID")
        secret_key = os.environ.get("S3_SECRET_ACCESS_KEY")
        addressing_style = os.environ.get("S3_ADDRESSING_STYLE", "auto")

        if not all([endpoint_url, access_key, secret_key]):
            logger.debug(
                "S3 environment variables not fully configured, skipping S3 direct download."
            )
            return None

        try:
            s3_config = BotoConfig(
                s3={"addressing_style": addressing_style},
                connect_timeout=5,
                read_timeout=15,
            )
            s3_client = boto3.client(
                "s3",
                endpoint_url=endpoint_url,
                aws_access_key_id=access_key,
                aws_secret_access_key=secret_key,
                config=s3_config,
            )
            response = s3_client.get_object(Bucket=bucket, Key=key)
            body = response["Body"]
            data = body.read(max_bytes + 1)
            body.close()
            if len(data) > max_bytes:
                return None
            return data
        except Exception as e:
            logger.warning(f"S3 direct download failed for {s3_path}: {e}")
            return None

    # ── Full 6-step image resolution (merged upstream + APA glob) ────────
    def _image_bytes_from_owui_file_id(
        self, file_id: str, max_bytes: int
    ) -> Optional[bytes]:
        if not file_id:
            return None
        if Files is None:
            logger.error(
                "Files model is not available (import failed). Cannot retrieve file content."
            )
            return None

        try:
            file_obj = Files.get_file_by_id(file_id)
        except Exception as e:
            logger.error(f"Files.get_file_by_id({file_id}) failed: {e}")
            return None
        if not file_obj:
            logger.warning(f"File {file_id} not found in database.")
            return None

        # 1. Try data field (DB stored)
        data_field = getattr(file_obj, "data", None)
        if isinstance(data_field, dict):
            blob_value = data_field.get("bytes")
            if isinstance(blob_value, (bytes, bytearray)):
                raw = bytes(blob_value)
                return raw if len(raw) <= max_bytes else None
            for key in ("b64", "base64", "data"):
                inline = data_field.get(key)
                if isinstance(inline, str) and inline.strip():
                    return self._decode_base64_limited(inline, max_bytes)

        # 2. Try S3 direct download (fastest for object storage)
        s3_path = getattr(file_obj, "path", None)
        if isinstance(s3_path, str) and s3_path.startswith("s3://"):
            s3_data = self._read_from_s3(s3_path, max_bytes)
            if s3_data is not None:
                return s3_data

        # 3. Try file paths (Disk stored) – multiple path variations
        for attr in ("path", "file_path", "absolute_path"):
            candidate = getattr(file_obj, attr, None)
            if isinstance(candidate, str) and candidate.strip():
                if re.match(r"^(s3://|gs://|https?://)", candidate, re.IGNORECASE):
                    logger.debug(f"Skipping local read for non-local path: {candidate}")
                    continue

                p = Path(candidate)

                # Attempt 1: As-is
                raw = self._read_file_bytes_limited(p, max_bytes)
                if raw is not None:
                    return raw

                # Attempt 2: Relative to ./data
                if not p.is_absolute():
                    try:
                        raw = self._read_file_bytes_limited(
                            Path("./data") / p, max_bytes
                        )
                        if raw is not None:
                            return raw
                    except Exception:
                        pass

                    # Attempt 3: Relative to /app/backend/data (Docker default)
                    try:
                        raw = self._read_file_bytes_limited(
                            Path("/app/backend/data") / p, max_bytes
                        )
                        if raw is not None:
                            return raw
                    except Exception:
                        pass

        # 3b. APA fork addition: glob-based UUID prefix search
        file_id_clean = file_id.split("/")[-1].strip()
        search_patterns = [
            f"/app/backend/data/uploads/{file_id_clean}*",
            f"/data/uploads/{file_id_clean}*",
        ]
        for pattern in search_patterns:
            matches = glob.glob(pattern)
            if matches:
                file_path = Path(matches[0])
                if file_path.exists() and file_path.is_file():
                    logger.info(f"Found image file via glob at: {file_path}")
                    raw = self._read_file_bytes_limited(file_path, max_bytes)
                    if raw is not None:
                        return raw

        # 4. Try URL (Object Storage / S3 Public URL)
        urls_to_try = []
        url_attr = getattr(file_obj, "url", None)
        if isinstance(url_attr, str) and url_attr:
            urls_to_try.append(url_attr)
        if isinstance(data_field, dict):
            url_data = data_field.get("url")
            if isinstance(url_data, str) and url_data:
                urls_to_try.append(url_data)

        if urls_to_try:
            import urllib.request

            for url in urls_to_try:
                if not url.startswith(("http://", "https://")):
                    continue
                try:
                    logger.info(
                        f"Attempting to download file {file_id} from URL: {url}"
                    )
                    req = urllib.request.Request(
                        url, headers={"User-Agent": "OpenWebUI-Export-Plugin"}
                    )
                    with urllib.request.urlopen(req, timeout=15) as response:
                        if 200 <= response.status < 300:
                            data = response.read(max_bytes + 1)
                            if len(data) <= max_bytes:
                                return data
                            else:
                                logger.warning(
                                    f"File {file_id} from URL is too large (> {max_bytes} bytes)"
                                )
                except Exception as e:
                    logger.warning(f"Failed to download {file_id} from {url}: {e}")

        # 5. Try fetching via Local API (upstream: authenticated internal fetch)
        if self._api_base_url:
            api_url = f"{self._api_base_url}/api/v1/files/{file_id}/content"
            try:
                import urllib.request

                headers = {"User-Agent": "OpenWebUI-Export-Plugin"}
                if self._api_token:
                    headers["Authorization"] = self._api_token
                req = urllib.request.Request(api_url, headers=headers)
                with urllib.request.urlopen(req, timeout=15) as response:
                    if 200 <= response.status < 300:
                        data = response.read(max_bytes + 1)
                        if len(data) <= max_bytes:
                            return data
            except Exception:
                pass

        # 6. Try direct content attributes (last ditch)
        for attr in ("content", "blob", "data"):
            raw = getattr(file_obj, attr, None)
            if isinstance(raw, (bytes, bytearray)):
                b = bytes(raw)
                return b if len(b) <= max_bytes else None

        logger.warning(
            f"File {file_id} found but no content accessible. Attributes: {dir(file_obj)}"
        )
        return None

    def _image_bytes_from_url(self, url: str, max_bytes: int) -> Optional[bytes]:
        """Download image from external HTTP/HTTPS URL."""
        if not isinstance(url, str) or not url.startswith(("http://", "https://")):
            return None
        try:
            import urllib.request

            req = urllib.request.Request(
                url, headers={"User-Agent": "OpenWebUI-Export-Plugin"}
            )
            with urllib.request.urlopen(req, timeout=15) as response:
                if 200 <= response.status < 300:
                    data = response.read(max_bytes + 1)
                    if len(data) <= max_bytes:
                        return data
        except Exception as e:
            logger.warning(f"Failed to download external image from {url}: {e}")
        return None

    def _add_image_placeholder(self, paragraph, alt: str, reason: str):
        label = (alt or "").strip() or "image"
        msg = f"[{label} not embedded: {reason}]"
        self._add_text_run(paragraph, msg, bold=False, italic=False, strike=False)

    def _try_embed_image(
        self, paragraph, image_bytes: bytes
    ) -> Tuple[bool, Optional[str]]:
        if not image_bytes:
            return False, "empty image bytes"
        try:
            run = paragraph.add_run()
            width = None
            if self._active_doc is not None:
                try:
                    width = self._available_block_width(self._active_doc)
                except Exception:
                    width = None
            run.add_picture(cast(Any, io.BytesIO(image_bytes)), width=width)
            return True, None
        except Exception as e:
            return False, str(e)

    def _embed_markdown_image(self, paragraph, alt: str, url: str):
        max_bytes = self._max_embed_image_bytes()
        u = (url or "").strip()

        if not u:
            self._add_image_placeholder(paragraph, alt, "missing URL")
            return

        image_bytes: Optional[bytes] = None

        # Strategy 1: Data URLs
        if u.lower().startswith("data:"):
            image_bytes = self._image_bytes_from_data_url(u, max_bytes)
            if image_bytes is None:
                self._add_image_placeholder(
                    paragraph,
                    alt,
                    f"invalid data URL or exceeds {self.valves.MAX_EMBED_IMAGE_MB}MB",
                )
                return

        # Strategy 2: OpenWebUI file IDs
        elif "/api/v1/files/" in u:
            file_id = self._extract_owui_api_file_id(u)
            if file_id:
                image_bytes = self._image_bytes_from_owui_file_id(file_id, max_bytes)
                if image_bytes is None:
                    self._add_image_placeholder(
                        paragraph, alt, f"file unavailable ({file_id})"
                    )
                    return
            else:
                self._add_image_placeholder(paragraph, alt, "invalid file ID")
                return

        # Strategy 3: External HTTP/HTTPS URLs
        elif u.lower().startswith(("http://", "https://")):
            image_bytes = self._image_bytes_from_url(u, max_bytes)
            if image_bytes is None:
                self._add_image_placeholder(
                    paragraph,
                    alt,
                    f"external image unavailable or exceeds {self.valves.MAX_EMBED_IMAGE_MB}MB",
                )
                return

        else:
            self._add_image_placeholder(paragraph, alt, f"unsupported URL format: {u}")
            return

        success, error_msg = self._try_embed_image(paragraph, image_bytes)
        if not success:
            self._add_image_placeholder(
                paragraph, alt, f"unsupported image type: {error_msg}"
            )

    # ── Markdown → DOCX core ────────────────────────────────────────────
    async def markdown_to_docx(
        self,
        markdown_text: str,
        top_heading: str = "",
        has_h1: bool = False,
        sources: Optional[List[dict]] = None,
        event_emitter: Optional[Callable] = None,
    ) -> Document:
        doc = Document()
        self._active_doc = doc
        try:
            self._mermaid_figure_counter = 0
            self._mermaid_placeholder_counter = 0
            self._caption_style_name = None
            self._citation_anchor_by_index = {}
            self._citation_refs = self._build_citation_refs(sources or [])
            self._bookmark_id_counter = 1
            for ref in self._citation_refs:
                self._citation_anchor_by_index[ref.idx] = ref.anchor

            self.set_document_default_font(doc)

            if top_heading and not has_h1:
                self.add_heading(doc, top_heading, 1)

            lines = markdown_text.split("\n")
            i = 0
            in_code_block = False
            code_block_content = []
            code_block_info_raw = ""
            code_block_lang = ""
            code_block_attrs: List[str] = []
            in_math_block = False
            math_block_delim = ""
            math_block_lines: List[str] = []
            in_list = False
            list_items = []
            list_type = None

            total_lines = len(lines)
            last_update_time = time.time()

            while i < len(lines):
                if event_emitter and time.time() - last_update_time > 2.0:
                    progress = int((i / total_lines) * 100)
                    await event_emitter(
                        {
                            "type": "status",
                            "data": {
                                "description": f"{self._get_msg('converting')} ({progress}%)",
                                "done": False,
                            },
                        }
                    )
                    last_update_time = time.time()

                line = lines[i]

                # Handle display math blocks
                if not in_code_block and self.valves.MATH_ENABLE:
                    single_line = self._extract_single_line_math(line)
                    if single_line is not None:
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                            in_list = False
                        self._add_display_equation(doc, single_line)
                        i += 1
                        continue

                    if not in_math_block:
                        stripped = line.strip()
                        if stripped in (r"\[", "$$"):
                            if in_list and list_items:
                                self.add_list_to_doc(doc, list_items, list_type)
                                list_items = []
                                in_list = False
                            in_math_block = True
                            math_block_delim = stripped
                            math_block_lines = []
                            i += 1
                            continue
                    else:
                        stripped = line.strip()
                        close = r"\]" if math_block_delim == r"\[" else "$$"
                        if stripped == close:
                            in_math_block = False
                            latex = "\n".join(math_block_lines).strip()
                            self._add_display_equation(doc, latex)
                            math_block_delim = ""
                            math_block_lines = []
                            i += 1
                            continue
                        math_block_lines.append(line)
                        i += 1
                        continue

                # Handle code blocks
                if line.strip().startswith("```"):
                    if not in_code_block:
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                            in_list = False
                        in_code_block = True
                        code_block_info_raw = line.strip()[3:].strip()
                        code_block_lang, code_block_attrs = self._parse_fence_info(
                            code_block_info_raw
                        )
                        code_block_content = []
                    else:
                        in_code_block = False
                        code_text = "\n".join(code_block_content)
                        if code_block_lang.lower() == "mermaid":
                            self._insert_mermaid_placeholder(doc, code_text)
                        else:
                            self.add_code_block(doc, code_text, code_block_lang)
                        code_block_content = []
                        code_block_info_raw = ""
                        code_block_lang = ""
                        code_block_attrs = []
                    i += 1
                    continue

                if in_code_block:
                    code_block_content.append(line)
                    i += 1
                    continue

                # Handle tables
                if line.strip().startswith("|") and line.strip().endswith("|"):
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    self.add_table(doc, table_lines)
                    continue

                # Handle headings
                header_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
                if header_match:
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    level = len(header_match.group(1))
                    text = header_match.group(2)
                    self.add_heading(doc, text, level)
                    i += 1
                    continue

                # Handle unordered lists
                unordered_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
                if unordered_match:
                    if not in_list or list_type != "unordered":
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                        in_list = True
                        list_type = "unordered"
                    indent = len(unordered_match.group(1)) // 2
                    list_items.append((indent, unordered_match.group(2)))
                    i += 1
                    continue

                # Handle ordered lists
                ordered_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
                if ordered_match:
                    if not in_list or list_type != "ordered":
                        if in_list and list_items:
                            self.add_list_to_doc(doc, list_items, list_type)
                            list_items = []
                        in_list = True
                        list_type = "ordered"
                    indent = len(ordered_match.group(1)) // 2
                    list_items.append((indent, ordered_match.group(2)))
                    i += 1
                    continue

                # Handle blockquotes
                if line.strip().startswith(">"):
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    blockquote_lines = []
                    while i < len(lines) and lines[i].strip().startswith(">"):
                        quote_line = re.sub(r"^>\s?", "", lines[i])
                        blockquote_lines.append(quote_line)
                        i += 1
                    self.add_blockquote(doc, "\n".join(blockquote_lines))
                    continue

                # Handle horizontal rules
                if re.match(r"^[-*_]{3,}$", line.strip()):
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    self.add_horizontal_rule(doc)
                    i += 1
                    continue

                # Handle empty lines
                if not line.strip():
                    if in_list and list_items:
                        self.add_list_to_doc(doc, list_items, list_type)
                        list_items = []
                        in_list = False
                    i += 1
                    continue

                # Handle normal paragraphs
                if in_list and list_items:
                    self.add_list_to_doc(doc, list_items, list_type)
                    list_items = []
                    in_list = False
                self.add_paragraph(doc, line)
                i += 1

            if in_list and list_items:
                self.add_list_to_doc(doc, list_items, list_type)

            if in_math_block and math_block_lines:
                self.add_paragraph(doc, r"\[")
                for l in math_block_lines:
                    self.add_paragraph(doc, l)
                self.add_paragraph(doc, r"\]")

            if self._citation_refs:
                self._add_references_section(doc)

            return doc
        finally:
            self._active_doc = None

    # ── math helpers ─────────────────────────────────────────────────────
    def _extract_single_line_math(self, line: str) -> Optional[str]:
        s = line.strip()
        m = re.match(r"^\\\[(.*)\\\]$", s)
        if m:
            return m.group(1).strip()
        m = re.match(r"^\$\$(.*)\$\$$", s)
        if m:
            return m.group(1).strip()
        return None

    def _strip_reasoning_blocks(self, text: str) -> str:
        if not text:
            return text
        cur = text
        for _ in range(10):
            prev = cur
            cur = _ALL_DETAILS_RE.sub("", cur)
            cur = _THINK_RE.sub("", cur)
            cur = _ANALYSIS_RE.sub("", cur)
            if cur == prev:
                break
        cur = re.sub(r"\n{4,}", "\n\n\n", cur)
        return cur

    def _add_display_equation(self, doc: Document, latex: str):
        latex = (latex or "").strip()
        if not latex:
            return
        if not LATEX_MATH_AVAILABLE:
            p = doc.add_paragraph()
            self._add_text_run(
                p, f"[Equation: {latex}]", bold=False, italic=True, strike=False
            )
            return
        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cast(Any, para)._p.append(self._wrap_omml_for_word(omml))
        except Exception as exc:
            logger.warning(f"Math conversion failed; falling back to text: {exc}")
            self.add_code_block(doc, latex, "latex")

    def _wrap_omml_for_word(self, omml: str):
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xml = f'<m:oMathPara xmlns:m="{m_ns}" xmlns:w="{w_ns}">{omml}</m:oMathPara>'
        return parse_xml(xml)

    def _add_inline_equation(
        self,
        paragraph,
        latex: str,
        bold: bool = False,
        italic: bool = False,
        strike: bool = False,
    ):
        latex = (latex or "").strip()
        if not latex:
            return
        if not self.valves.MATH_ENABLE or not LATEX_MATH_AVAILABLE:
            self._add_text_run(
                paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike
            )
            return
        try:
            mathml = latex_to_mathml(latex)
            omml = mathml2omml.convert(mathml)
            o_math = self._omml_oMath_element(omml)
            run = paragraph.add_run()
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
            cast(Any, run)._r.append(o_math)
        except Exception as exc:
            logger.warning(f"Inline math conversion failed; keeping literal: {exc}")
            self._add_text_run(
                paragraph, f"\\({latex}\\)", bold=bold, italic=italic, strike=strike
            )

    def _omml_oMath_element(self, omml: str):
        m_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        s = (omml or "").strip()
        if s.startswith("<m:oMath>") and s.endswith("</m:oMath>"):
            inner = s[len("<m:oMath>") : -len("</m:oMath>")]
            s = f'<m:oMath xmlns:m="{m_ns}">{inner}</m:oMath>'
        elif s.startswith("<m:oMath") and "xmlns:m=" not in s.split(">", 1)[0]:
            s = s.replace("<m:oMath", f'<m:oMath xmlns:m="{m_ns}"', 1)
        return parse_xml(s)

    # ── citations ────────────────────────────────────────────────────────
    def _build_citation_refs(self, sources: List[dict]) -> List[_CitationRef]:
        citation_idx_map: Dict[str, int] = {}
        refs_by_idx: Dict[int, _CitationRef] = {}

        for source in sources or []:
            if not isinstance(source, dict):
                continue
            documents = source.get("document") or []
            metadatas = source.get("metadata") or []
            src_info = source.get("source") or {}
            src_name = src_info.get("name") if isinstance(src_info, dict) else None
            src_id_default = src_info.get("id") if isinstance(src_info, dict) else None
            src_urls = src_info.get("urls") if isinstance(src_info, dict) else None
            if not isinstance(documents, list):
                documents = []
            if not isinstance(metadatas, list):
                metadatas = []

            for idx_doc, _doc_text in enumerate(documents):
                meta = metadatas[idx_doc] if idx_doc < len(metadatas) else {}
                if not isinstance(meta, dict):
                    meta = {}
                source_id = meta.get("source") or src_id_default or "N/A"
                source_id_str = str(source_id)
                if source_id_str not in citation_idx_map:
                    citation_idx_map[source_id_str] = len(citation_idx_map) + 1
                idx = citation_idx_map[source_id_str]
                if idx in refs_by_idx:
                    continue

                url: Optional[str] = None
                if isinstance(source_id, str) and re.match(r"^https?://", source_id):
                    url = source_id
                elif isinstance(meta.get("url"), str) and re.match(
                    r"^https?://", meta["url"]
                ):
                    url = meta["url"]
                elif isinstance(src_urls, list) and src_urls:
                    if isinstance(src_urls[0], str) and re.match(
                        r"^https?://", src_urls[0]
                    ):
                        url = src_urls[0]

                title = (
                    (meta.get("title") if isinstance(meta.get("title"), str) else None)
                    or (meta.get("name") if isinstance(meta.get("name"), str) else None)
                    or (
                        src_name
                        if isinstance(src_name, str) and src_name.strip()
                        else None
                    )
                    or (url if url else None)
                    or source_id_str
                )

                anchor = f"OWUIRef{idx}"
                refs_by_idx[idx] = _CitationRef(
                    idx=idx,
                    anchor=anchor,
                    title=title,
                    url=url,
                    source_id=source_id_str,
                )

        return [refs_by_idx[i] for i in sorted(refs_by_idx.keys())]

    def _add_bookmark(self, paragraph, name: str):
        bookmark_id = self._bookmark_id_counter
        self._bookmark_id_counter += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        p = cast(Any, paragraph)._p
        p.insert(0, start)
        p.append(end)

    def _add_internal_hyperlink(self, paragraph, display_text: str, anchor: str):
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)
        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    def _add_references_section(self, doc: Document):
        self.add_heading(doc, self._get_msg("references"), 2)
        for ref in self._citation_refs:
            para = doc.add_paragraph(style="List Number")
            self._add_bookmark(para, ref.anchor)
            if ref.url:
                self._add_hyperlink(para, ref.title, ref.url, display_text=ref.title)
            else:
                self._add_text_run(
                    para, ref.title, bold=False, italic=False, strike=False
                )

    # ── Mermaid helpers ──────────────────────────────────────────────────
    def _parse_fence_info(self, info_raw: str) -> Tuple[str, List[str]]:
        parts = [p for p in (info_raw or "").split() if p.strip()]
        if not parts:
            return "", []
        return parts[0], parts[1:]

    def _normalize_mermaid_text(self, source: str) -> str:
        text = (source or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        return text + "\n"

    def _prepare_mermaid_for_js(self, source: str) -> str:
        return self._strip_mermaid_title_for_render(source)

    def _png_with_text_chunk(self, png_bytes: bytes, keyword: str, value: str) -> bytes:
        if not png_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
            return png_bytes
        keyword_b = (keyword or "owui").encode("latin-1", errors="ignore")[:79]
        keyword_b = keyword_b.replace(b"\x00", b"") or b"owui"
        value_b = (value or "").encode("latin-1", errors="ignore")
        data = keyword_b + b"\x00" + value_b
        chunk_type = b"tEXt"
        crc = zlib.crc32(chunk_type + data) & 0xFFFFFFFF
        chunk = (
            struct.pack("!I", len(data)) + chunk_type + data + struct.pack("!I", crc)
        )
        out = bytearray()
        out.extend(png_bytes[:8])
        offset = 8
        inserted = False
        while offset + 8 <= len(png_bytes):
            length = struct.unpack("!I", png_bytes[offset : offset + 4])[0]
            ctype = png_bytes[offset + 4 : offset + 8]
            chunk_total = 12 + length
            if offset + chunk_total > len(png_bytes):
                break
            if ctype == b"IEND" and not inserted:
                out.extend(chunk)
                inserted = True
            out.extend(png_bytes[offset : offset + chunk_total])
            offset += chunk_total
            if ctype == b"IEND":
                break
        if not inserted:
            return png_bytes
        return bytes(out)

    def _make_mermaid_placeholder_png(self, seed: str) -> bytes:
        return self._png_with_text_chunk(_TRANSPARENT_1PX_PNG, "owui", seed)

    def _dummy_mermaid_svg_bytes(self) -> bytes:
        return '<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1" viewBox="0 0 1 1"></svg>'.encode(
            "utf-8"
        )

    def _insert_mermaid_placeholder(self, doc: Document, mermaid_source: str):
        caption_title: Optional[str] = (
            self._extract_mermaid_title(mermaid_source)
            if self.valves.MERMAID_CAPTIONS_ENABLE
            else None
        )
        source_for_render = mermaid_source
        if self.valves.MERMAID_OPTIMIZE_LAYOUT:
            source_for_render = re.sub(
                r"^(graph|flowchart)\s+LR\b",
                r"\1 TD",
                source_for_render,
                flags=re.MULTILINE | re.IGNORECASE,
            )
        source_for_render = self._prepare_mermaid_for_js(source_for_render)

        self._mermaid_placeholder_counter += 1
        seed = hashlib.sha256(
            f"{self._mermaid_placeholder_counter}\n{source_for_render}".encode(
                "utf-8", errors="replace"
            )
        ).hexdigest()[:16]
        png_bytes = self._make_mermaid_placeholder_png(seed)

        try:
            shape = doc.add_picture(cast(Any, io.BytesIO(png_bytes)))
        except Exception as e:
            logger.warning(f"Failed to add Mermaid placeholder image: {e}")
            self.add_paragraph(doc, f"[Mermaid placeholder failed: {e}]")
            return
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        self._attach_svg_blip(doc, shape, self._dummy_mermaid_svg_bytes())

        try:
            encoded = quote(source_for_render)
            inline = shape._inline
            docPr = inline.docPr
            docPr.set("descr", f"MERMAID_SRC:{encoded}")
            docPr.set("title", "Mermaid Diagram Placeholder")
        except Exception as exc:
            logger.warning(f"Failed to annotate Mermaid placeholder: {exc}")

        self._add_mermaid_caption(doc, caption_title)

    def _extract_mermaid_title(self, source: str) -> Optional[str]:
        lines = self._normalize_mermaid_text(source).split("\n")
        header_found = False
        for raw in lines:
            line = raw.strip()
            if not line:
                continue
            if line.startswith("%%{") and line.endswith("}%%"):
                continue
            if line.startswith("%%"):
                continue
            if not header_found:
                header_found = True
                mt = re.match(
                    r"^(?P<header>\S.*?)(?:\s+title\s*:?\s+)(?P<title>.+)$",
                    line,
                    re.IGNORECASE,
                )
                if mt:
                    title = (mt.group("title") or "").strip().strip('"').strip("'")
                    if title:
                        return title
                continue
            m = re.match(r'^title\s*:?\s+"(.+)"\s*$', line, re.IGNORECASE)
            if m:
                return m.group(1).strip()
            m = re.match(r"^title\s*:?\s+(.+)$", line, re.IGNORECASE)
            if m:
                return m.group(1).strip().strip('"').strip("'")
        return None

    def _strip_mermaid_title_for_render(self, source: str) -> str:
        lines = self._normalize_mermaid_text(source).split("\n")
        out: List[str] = []
        header_found = False
        title_stripped = False
        meaningful_after_header = False
        for raw in lines:
            line = raw.rstrip("\n")
            stripped = line.strip()
            if not stripped:
                out.append(line)
                continue
            if stripped.startswith("%%{") and stripped.endswith("}%%"):
                out.append(line)
                continue
            if stripped.startswith("%%"):
                out.append(line)
                continue
            if not header_found:
                header_found = True
                mt = re.match(
                    r"^(?P<header>\S.*?)(?:\s+title\s*:?\s+)(?P<title>.+)$",
                    stripped,
                    re.IGNORECASE,
                )
                if mt:
                    cleaned = (mt.group("header") or "").strip()
                    out.append(cleaned if cleaned else stripped)
                    title_stripped = True
                    continue
                out.append(line)
                continue
            if not title_stripped and not meaningful_after_header:
                if re.match(r'^title\s*:?\s+(".+"|.+)$', stripped, re.IGNORECASE):
                    title_stripped = True
                    continue
            meaningful_after_header = True
            out.append(line)
        return "\n".join(out).strip() + "\n"

    def _ensure_caption_style(self, doc: Document) -> str:
        if self._caption_style_name is not None:
            return self._caption_style_name
        style_name = (self.valves.MERMAID_CAPTION_STYLE or "").strip()
        if style_name == "":
            self._caption_style_name = ""
            return ""
        try:
            _ = doc.styles[style_name]
            self._caption_style_name = style_name
            return style_name
        except KeyError:
            pass
        if style_name.lower() == "caption":
            style_name = "OWUI Caption"
        try:
            _ = doc.styles[style_name]
            self._caption_style_name = style_name
            return style_name
        except KeyError:
            pass
        try:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = self.valves.FONT_LATIN
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(80, 80, 80)
            style.paragraph_format.space_before = Pt(2)
            style.paragraph_format.space_after = Pt(8)
            self._caption_style_name = style_name
            return style_name
        except Exception:
            self._caption_style_name = "Normal"
            return "Normal"

    def _add_mermaid_caption(self, doc: Document, title: Optional[str]):
        if not self.valves.MERMAID_CAPTIONS_ENABLE:
            return
        prefix = (self.valves.MERMAID_CAPTION_PREFIX or "").strip()
        if prefix == "":
            prefix = self._get_msg("figure_prefix")
        if prefix == "" and not title:
            return
        self._mermaid_figure_counter += 1
        if prefix == "":
            caption = title or ""
        else:
            base = f"{prefix} {self._mermaid_figure_counter}"
            caption = f"{base}: {title}" if title else base
        if caption == "":
            return
        para = doc.add_paragraph()
        style_name = self._ensure_caption_style(doc)
        if style_name:
            para.style = style_name
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_formatted_text(para, caption)

    def _available_block_width(self, doc: Document):
        section = doc.sections[0]
        return section.page_width - section.left_margin - section.right_margin

    def _attach_svg_blip(self, doc: Document, inline_shape: Any, svg_bytes: bytes):
        if not svg_bytes:
            return
        try:
            pkg = doc.part.package
            partname = pkg.next_partname("/word/media/image%d.svg")
            from docx.opc.part import Part

            svg_part = Part(partname, "image/svg+xml", svg_bytes)
            rid_svg = doc.part.relate_to(svg_part, RT.IMAGE)
            inline = inline_shape._inline
            blips = inline.xpath(".//a:blip")
            if not blips:
                return
            blip = blips[0]
            existing = blip.xpath(".//asvg:svgBlip")
            if existing:
                existing[0].set(qn("r:embed"), rid_svg)
                return
            extLst = OxmlElement("a:extLst")
            ext = OxmlElement("a:ext")
            ext.set("uri", "{96DAC541-7B7A-43D3-8B79-37D633B846F1}")
            svgBlip = OxmlElement("asvg:svgBlip")
            svgBlip.set(qn("r:embed"), rid_svg)
            ext.append(svgBlip)
            extLst.append(ext)
            blip.append(extLst)
        except Exception as exc:
            logger.warning(f"Failed to attach SVG blip; keeping PNG fallback: {exc}")

    # ── document font setup (APA + upstream East Asian fallback) ─────────
    def set_document_default_font(self, doc: Document):
        """Set document default font and APA page layout (margins, indent, spacing)."""
        # ── APA Page Margins: 1 inch on all sides ────────────────────────
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # ── Normal style: font + paragraph formatting ────────────────────
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.valves.FONT_LATIN
        font.size = Pt(12)  # APA Standard size

        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()

        rFonts.set(qn("w:ascii"), self.valves.FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), self.valves.FONT_LATIN)  # Critical for Greek
        rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
        rFonts.set(qn("w:cs"), self.valves.FONT_LATIN)  # Complex Script

        # APA: Double spacing, no extra space after, 0.5" first-line indent
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Inches(0.5)

        # ── APA Page Numbers: top-right ──────────────────────────────────
        self._add_page_numbers(doc)

    def _add_page_numbers(self, doc: Document):
        """Add APA-compliant page numbers (top-right) to document header."""
        try:
            section = doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False

            # Clear any existing header content
            for para in header.paragraphs:
                para.clear()

            # Use the first paragraph or create one
            if header.paragraphs:
                para = header.paragraphs[0]
            else:
                para = header.add_paragraph()

            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

            # Insert PAGE field code via XML
            run = para.add_run()
            run.font.name = self.valves.FONT_LATIN
            run.font.size = Pt(12)

            fld_char_begin = OxmlElement("w:fldChar")
            fld_char_begin.set(qn("w:fldCharType"), "begin")
            run._element.append(fld_char_begin)

            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = " PAGE "
            run._element.append(instr_text)

            fld_char_end = OxmlElement("w:fldChar")
            fld_char_end.set(qn("w:fldCharType"), "end")
            run._element.append(fld_char_end)
        except Exception as e:
            logger.warning(f"Failed to add page numbers: {e}")

    # ── APA heading formatting ───────────────────────────────────────────
    def add_heading(self, doc: Document, text: str, level: int):
        """Add APA-compliant heading using Word's built-in styles (for TOC support)."""
        heading_level = min(level, 5)  # APA uses 5 levels max

        heading = doc.add_heading(level=heading_level)
        heading.text = ""

        run = heading.add_run(text)
        run.font.name = self.valves.FONT_LATIN
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Upstream: set East Asian font on heading runs too
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
        rFonts.set(qn("w:hAnsi"), self.valves.FONT_LATIN)

        if heading_level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
        elif heading_level == 2:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.bold = True
        elif heading_level == 3:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.bold = True
            run.italic = True
        elif heading_level == 4:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.paragraph_format.left_indent = Inches(0.5)
            run.bold = True
        else:  # 5
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.paragraph_format.left_indent = Inches(0.5)
            run.bold = True
            run.italic = True

        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.first_line_indent = Inches(0)  # Headings: no indent

    def add_paragraph(self, doc: Document, text: str):
        paragraph = doc.add_paragraph()
        self.add_formatted_text(paragraph, text)

    def add_formatted_text(self, paragraph, text: str):
        self._add_inline_segments(
            paragraph, text or "", bold=False, italic=False, strike=False
        )

    # ── text run (with upstream East Asian font fallback) ────────────────
    def _add_text_run(self, paragraph, s: str, bold: bool, italic: bool, strike: bool):
        if not s:
            return
        run = paragraph.add_run(s)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True

        # Ported from upstream: Explicitly set East Asian font on every run
        # to prevent MS Gothic fallback for any CJK content
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)

    def _add_inline_code(self, paragraph, s: str):
        if s == "":
            return

        def _add_code_run(chunk: str):
            if not chunk:
                return
            run = paragraph.add_run(chunk)
            run.font.name = self.valves.FONT_CODE
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), self.valves.FONT_CODE)
            run.font.size = Pt(11)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E8E8E8")
            run._element.rPr.append(shading)

        i = 0
        for m in _AUTO_URL_RE.finditer(s):
            start, end = m.span()
            if start > i:
                _add_code_run(s[i:start])
            raw = m.group(0)
            trimmed = raw
            while trimmed and trimmed[-1] in ".,;:!?)]}":
                trimmed = trimmed[:-1]
            suffix = raw[len(trimmed) :]
            normalized = self._normalize_url(trimmed)
            if normalized:
                self._add_hyperlink_code(
                    paragraph, display_text=trimmed, url=normalized
                )
            else:
                _add_code_run(raw)
            if suffix:
                _add_code_run(suffix)
            i = end
        if i < len(s):
            _add_code_run(s[i:])

    def _add_hyperlink_code(self, paragraph, display_text: str, url: str):
        u = self._normalize_url(url)
        if not u:
            self._add_inline_code(paragraph, display_text)
            return
        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            self._add_inline_code(paragraph, display_text)
            return

        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self.valves.FONT_CODE)
        rFonts.set(qn("w:hAnsi"), self.valves.FONT_CODE)
        rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "22")
        rPr.append(sz_cs)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8E8E8")
        rPr.append(shading)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text
        new_run.append(t)
        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    # ── inline segment parser ────────────────────────────────────────────
    def _add_inline_segments(
        self, paragraph, text: str, bold: bool, italic: bool, strike: bool
    ):
        i = 0
        n = len(text)

        def next_special(start: int) -> int:
            candidates = []
            for ch in ("`", "!", "[", "*", "_", "~", "$", "\\"):
                idx = text.find(ch, start)
                if idx != -1:
                    candidates.append(idx)
            idx = text.find(r"\(", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("http://", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("https://", start)
            if idx != -1:
                candidates.append(idx)
            idx = text.find("www.", start)
            if idx != -1:
                candidates.append(idx)
            return min(candidates) if candidates else n

        while i < n:
            # Markdown image: ![alt](url)
            if text.startswith("![", i):
                close = text.find("]", i + 2)
                if close != -1 and close + 1 < n and text[close + 1] == "(":
                    close_paren = text.find(")", close + 2)
                    if close_paren != -1:
                        alt = text[i + 2 : close]
                        url = text[close + 2 : close_paren].strip()
                        if url.startswith("<") and url.endswith(">") and len(url) >= 2:
                            url = url[1:-1].strip()
                        self._embed_markdown_image(paragraph, alt=alt, url=url)
                        i = close_paren + 1
                        continue

            if text[i] == "`":
                j = text.find("`", i + 1)
                if j != -1:
                    self._add_inline_code(paragraph, text[i + 1 : j])
                    i = j + 1
                    continue

            if text.startswith(r"\(", i):
                j = text.find(r"\)", i + 2)
                if j != -1:
                    self._add_inline_equation(
                        paragraph,
                        text[i + 2 : j],
                        bold=bold,
                        italic=italic,
                        strike=strike,
                    )
                    i = j + 2
                    continue

            if text[i] == "\\":
                if i + 1 < n:
                    ch = text[i + 1]
                    if ch in "\\`*_{}[]()#+-.!|$":
                        self._add_text_run(paragraph, ch, bold, italic, strike)
                        i += 2
                        continue
                self._add_text_run(paragraph, "\\", bold, italic, strike)
                i += 1
                continue

            if text[i] == "_":
                run_len = 0
                while i + run_len < n and text[i + run_len] == "_":
                    run_len += 1
                if run_len >= 4:
                    self._add_text_run(
                        paragraph, text[i : i + run_len], bold, italic, strike
                    )
                    i += run_len
                    continue

            if text[i] == "*":
                run_len = 0
                while i + run_len < n and text[i + run_len] == "*":
                    run_len += 1
                if run_len >= 4:
                    self._add_text_run(
                        paragraph, text[i : i + run_len], bold, italic, strike
                    )
                    i += run_len
                    continue

            if text[i] == "~":
                run_len = 0
                while i + run_len < n and text[i + run_len] == "~":
                    run_len += 1
                if run_len >= 4:
                    self._add_text_run(
                        paragraph, text[i : i + run_len], bold, italic, strike
                    )
                    i += run_len
                    continue

            # Inline $…$ math
            if (
                text[i] == "$"
                and self.valves.MATH_ENABLE
                and self.valves.MATH_INLINE_DOLLAR_ENABLE
            ):
                if text.startswith("$$", i):
                    self._add_text_run(paragraph, "$", bold, italic, strike)
                    i += 1
                    continue
                if i + 1 >= n or text[i + 1].isspace():
                    self._add_text_run(paragraph, "$", bold, italic, strike)
                    i += 1
                    continue
                if i > 0 and text[i - 1].isalnum():
                    self._add_text_run(paragraph, "$", bold, italic, strike)
                    i += 1
                    continue

                j = i + 1
                while True:
                    j = text.find("$", j)
                    if j == -1:
                        break
                    if j > 0 and text[j - 1] == "\\":
                        j += 1
                        continue
                    break

                if j != -1:
                    inner = text[i + 1 : j]
                    if (
                        inner
                        and "\n" not in inner
                        and not inner[0].isspace()
                        and not inner[-1].isspace()
                    ):
                        if _CURRENCY_NUMBER_RE.match(inner) and (
                            i == 0 or text[i - 1].isspace()
                        ):
                            self._add_text_run(paragraph, "$", bold, italic, strike)
                            i += 1
                            continue
                        if j + 1 < n and text[j + 1].isdigit():
                            self._add_text_run(paragraph, "$", bold, italic, strike)
                            i += 1
                            continue
                        self._add_inline_equation(
                            paragraph, inner, bold=bold, italic=italic, strike=strike
                        )
                        i = j + 1
                        continue

                self._add_text_run(paragraph, "$", bold, italic, strike)
                i += 1
                continue

            if text.startswith("~~", i):
                j = text.find("~~", i + 2)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 2 : j],
                        bold=bold,
                        italic=italic,
                        strike=True,
                    )
                    i = j + 2
                    continue

            if text.startswith("**", i):
                j = text.find("**", i + 2)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 2 : j],
                        bold=True,
                        italic=italic,
                        strike=strike,
                    )
                    i = j + 2
                    continue

            if text.startswith("__", i):
                j = text.find("__", i + 2)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 2 : j],
                        bold=True,
                        italic=italic,
                        strike=strike,
                    )
                    i = j + 2
                    continue

            if text[i] == "*" and (i + 1 >= n or text[i + 1] != "*"):
                j = text.find("*", i + 1)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 1 : j],
                        bold=bold,
                        italic=True,
                        strike=strike,
                    )
                    i = j + 1
                    continue

            if text[i] == "_" and (i + 1 >= n or text[i + 1] != "_"):
                j = text.find("_", i + 1)
                if j != -1:
                    self._add_inline_segments(
                        paragraph,
                        text[i + 1 : j],
                        bold=bold,
                        italic=True,
                        strike=strike,
                    )
                    i = j + 1
                    continue

            if text[i] == "[":
                close = text.find("]", i + 1)
                if close != -1 and close + 1 < n and text[close + 1] == "(":
                    close_paren = text.find(")", close + 2)
                    if close_paren != -1:
                        label = text[i + 1 : close]
                        url = text[close + 2 : close_paren]
                        self._add_hyperlink(paragraph, label, url)
                        i = close_paren + 1
                        continue
                if close != -1:
                    inner = text[i + 1 : close].strip()
                    if inner.isdigit():
                        idx = int(inner)
                        anchor = self._citation_anchor_by_index.get(idx)
                        if anchor:
                            self._add_internal_hyperlink(paragraph, f"[{idx}]", anchor)
                            i = close + 1
                            continue

            m = _AUTO_URL_RE.match(text, i)
            if m:
                raw = m.group(0)
                trimmed = raw
                while trimmed and trimmed[-1] in ".,;:!?)]}":
                    trimmed = trimmed[:-1]
                suffix = raw[len(trimmed) :]
                normalized = self._normalize_url(trimmed)
                if normalized:
                    self._add_hyperlink(
                        paragraph, trimmed, normalized, display_text=trimmed
                    )
                else:
                    self._add_text_run(paragraph, raw, bold, italic, strike)
                    i += len(raw)
                    continue
                if suffix:
                    self._add_text_run(paragraph, suffix, bold, italic, strike)
                i += len(raw)
                continue

            j = next_special(i)
            if j == i:
                self._add_text_run(paragraph, text[i], bold, italic, strike)
                i += 1
            else:
                self._add_text_run(paragraph, text[i:j], bold, italic, strike)
                i = j

    def _normalize_url(self, url: str) -> str:
        u = (url or "").strip()
        if u.lower().startswith("www."):
            u = "https://" + u
        while u and u[-1] in ".,;:!?)]}":
            u = u[:-1]
        return u

    def _add_hyperlink(
        self, paragraph, text: str, url: str, display_text: Optional[str] = None
    ):
        u = self._normalize_url(url)
        if not u:
            run = paragraph.add_run(display_text or text)
            # Upstream: set East Asian font on fallback runs
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
            return

        part = getattr(paragraph, "part", None)
        if part is None or not hasattr(part, "relate_to"):
            run = paragraph.add_run(display_text or text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
            return

        r_id = part.relate_to(u, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        # Upstream: East Asian font on hyperlink runs
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
        rPr.append(rFonts)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)
        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rPr.append(u_el)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = display_text or text
        new_run.append(t)
        hyperlink.append(new_run)
        cast(Any, paragraph)._p.append(hyperlink)

    # ── code blocks ──────────────────────────────────────────────────────
    def add_code_block(self, doc: Document, code: str, language: str = ""):
        TOKEN_COLORS = {
            Token.Keyword: RGBColor(0, 92, 197),
            Token.Keyword.Constant: RGBColor(0, 92, 197),
            Token.Keyword.Declaration: RGBColor(0, 92, 197),
            Token.Keyword.Namespace: RGBColor(0, 92, 197),
            Token.Keyword.Type: RGBColor(0, 92, 197),
            Token.Name.Function: RGBColor(0, 0, 0),
            Token.Name.Class: RGBColor(38, 82, 120),
            Token.Name.Decorator: RGBColor(170, 51, 0),
            Token.Name.Builtin: RGBColor(0, 110, 71),
            Token.String: RGBColor(196, 26, 22),
            Token.String.Doc: RGBColor(109, 120, 133),
            Token.Comment: RGBColor(109, 120, 133),
            Token.Comment.Single: RGBColor(109, 120, 133),
            Token.Comment.Multiline: RGBColor(109, 120, 133),
            Token.Number: RGBColor(28, 0, 207),
            Token.Number.Integer: RGBColor(28, 0, 207),
            Token.Number.Float: RGBColor(28, 0, 207),
            Token.Operator: RGBColor(90, 99, 120),
            Token.Punctuation: RGBColor(0, 0, 0),
        }

        def get_token_color(token_type):
            while token_type:
                if token_type in TOKEN_COLORS:
                    return TOKEN_COLORS[token_type]
                token_type = token_type.parent
            return None

        if language:
            lang_para = doc.add_paragraph()
            lang_para.paragraph_format.space_before = Pt(6)
            lang_para.paragraph_format.space_after = Pt(0)
            lang_para.paragraph_format.left_indent = Cm(0.5)
            lang_para.paragraph_format.first_line_indent = Inches(0)
            lang_run = lang_para.add_run(language.upper())
            lang_run.font.name = self.valves.FONT_CODE
            lang_run.font.size = Pt(9)
            lang_run.font.color.rgb = RGBColor(100, 100, 100)
            lang_run.font.bold = True

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.space_before = Pt(3) if language else Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F7F7F7")
        paragraph._element.pPr.append(shading)

        if PYGMENTS_AVAILABLE and language:
            try:
                lexer = get_lexer_by_name(language, stripall=False)
            except Exception:
                lexer = TextLexer()

            tokens = list(lex(code, lexer))
            for token_type, token_value in tokens:
                if not token_value:
                    continue
                run = paragraph.add_run(token_value)
                run.font.name = self.valves.FONT_CODE
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
                run._element.rPr.rFonts.set(qn("w:hAnsi"), self.valves.FONT_CODE)
                run.font.size = Pt(10)
                color = get_token_color(token_type)
                if color:
                    run.font.color.rgb = color
                if token_type in Token.Keyword:
                    run.font.bold = True
        else:
            run = paragraph.add_run(code)
            run.font.name = self.valves.FONT_CODE
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.valves.FONT_CODE)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), self.valves.FONT_CODE)
            run.font.size = Pt(10)

    # ── tables ───────────────────────────────────────────────────────────
    def add_table(self, doc: Document, table_lines: List[str]):
        if len(table_lines) < 2:
            return

        def _validate_hex(c: str, default: str) -> str:
            c = c.strip().lstrip("#")
            if re.fullmatch(r"[0-9A-Fa-f]{6}", c):
                return c
            return default

        header_fill = _validate_hex(self.valves.TABLE_HEADER_COLOR, "E8E8E8")
        zebra_fill = _validate_hex(self.valves.TABLE_ZEBRA_COLOR, "F8F8F8")

        def _split_row(line: str) -> List[str]:
            raw = line.strip().strip("|")
            return [c.strip() for c in raw.split("|")]

        def _is_separator_row(cells: List[str]) -> bool:
            if not cells:
                return False
            ok = 0
            for c in cells:
                c = c.strip()
                if re.fullmatch(r":?-{3,}:?", c):
                    ok += 1
            return ok == len(cells)

        def _col_align(cell: str) -> WD_ALIGN_PARAGRAPH:
            s = (cell or "").strip()
            if s.startswith(":") and s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.CENTER
            if s.endswith(":"):
                return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT

        def _set_cell_shading(cell, fill: str):
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        raw_rows = [_split_row(l) for l in table_lines if l.strip().startswith("|")]
        if not raw_rows:
            return

        sep_idx = 1 if len(raw_rows) > 1 and _is_separator_row(raw_rows[1]) else -1
        header = raw_rows[0]
        body = raw_rows[sep_idx + 1 :] if sep_idx >= 0 else raw_rows[1:]

        num_cols = max(len(header), *(len(r) for r in body)) if body else len(header)
        header = header + [""] * (num_cols - len(header))
        body = [r + [""] * (num_cols - len(r)) for r in body]

        aligns = [
            _col_align(c) for c in (raw_rows[1] if sep_idx == 1 else [""] * num_cols)
        ]

        table = doc.add_table(rows=1 + len(body), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cast(Any, table).autofit = False

        self._set_table_cell_margins(table, top=60, bottom=60, left=90, right=90)

        available_width = int(self._available_block_width(doc))
        min_col = max(int(Inches(0.55)), available_width // max(1, num_cols * 3))

        def _plain_len(s: str) -> int:
            t = re.sub(r"`([^`]+)`", r"\1", s or "")
            t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", t)
            t = re.sub(r"\s+", " ", t).strip()
            return len(t)

        weights: List[int] = []
        for ci in range(num_cols):
            max_len = _plain_len(header[ci])
            for r in body:
                max_len = max(max_len, _plain_len(r[ci]))
            weights.append(max(1, min(max_len, 40)))

        sum_w = sum(weights) or 1
        widths = [max(min_col, int(available_width * w / sum_w)) for w in weights]
        total = sum(widths)
        if total > available_width:
            even = max(1, available_width // max(1, num_cols))
            widths = [even] * num_cols
            total = sum(widths)
        if total < available_width:
            rem = available_width - total
            order = sorted(range(num_cols), key=lambda i: weights[i], reverse=True)
            oi = 0
            while rem > 0 and order:
                widths[order[oi % len(order)]] += 1
                rem -= 1
                oi += 1

        for ci, w in enumerate(widths):
            table.columns[ci].width = w
            for row in table.rows:
                row.cells[ci].width = w

        def _format_cell_paragraph(para, align: WD_ALIGN_PARAGRAPH):
            para.alignment = align
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.first_line_indent = Inches(0)

        def _fill_cell(cell, text: str, align: WD_ALIGN_PARAGRAPH, bold: bool = False):
            cell.text = ""
            parts = [
                p for p in re.split(r"(?:<br\s*/?>|\n)", text or "") if p is not None
            ]
            if not parts:
                parts = [""]
            for pi, part in enumerate(parts):
                para = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                _format_cell_paragraph(para, align)
                self.add_formatted_text(para, part)
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = self.valves.FONT_LATIN
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn("w:eastAsia"), self.valves.FONT_ASIAN)
                    rFonts.set(qn("w:hAnsi"), self.valves.FONT_LATIN)
                    if bold:
                        run.bold = True

        header_row = table.rows[0]
        self._set_table_header_row_repeat(header_row)
        for ci in range(num_cols):
            cell = header_row.cells[ci]
            _set_cell_shading(cell, header_fill)
            _fill_cell(
                cell,
                header[ci],
                aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
            )

        for ri, row_data in enumerate(body, start=1):
            row = table.rows[ri]
            for ci in range(num_cols):
                cell = row.cells[ci]
                if (ri % 2) == 0:
                    _set_cell_shading(cell, zebra_fill)
                _fill_cell(
                    cell,
                    row_data[ci],
                    aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT,
                )

    def _set_table_cell_margins(
        self, table, top: int, bottom: int, left: int, right: int
    ):
        tbl_pr = cast(Any, table)._tbl.tblPr
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        for tag, val in (
            ("top", top),
            ("bottom", bottom),
            ("left", left),
            ("right", right),
        ):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")
            tbl_cell_mar.append(el)
        tbl_pr.append(tbl_cell_mar)

    def _set_table_header_row_repeat(self, row):
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    # ── lists ────────────────────────────────────────────────────────────
    def add_list_to_doc(
        self, doc: Document, items: List[Tuple[int, str]], list_type: str
    ):
        for indent, text in items:
            paragraph = doc.add_paragraph()
            if list_type == "unordered":
                paragraph.style = "List Bullet"
            else:
                paragraph.style = "List Number"
            paragraph.paragraph_format.left_indent = Cm(0.5 * (indent + 1))
            paragraph.paragraph_format.first_line_indent = Inches(0)
            self.add_formatted_text(paragraph, text)

    # ── horizontal rule ──────────────────────────────────────────────────
    def add_horizontal_rule(self, doc: Document):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── blockquote ───────────────────────────────────────────────────────
    def add_blockquote(self, doc: Document, text: str):
        for line in text.split("\n"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "CCCCCC")
            pBdr.append(left)
            pPr.append(pBdr)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F9F9F9")
            pPr.append(shading)
            self.add_formatted_text(paragraph, line)
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(85, 85, 85)
                run.italic = True
